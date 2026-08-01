from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor
import os

BASE = os.path.dirname(os.path.dirname(__file__))
FRONT_MD = os.path.join(BASE, 'FRONT_MATTER_PAGES.md')
TEMPLATE_MD = os.path.join(BASE, 'MINI_PROJECT_REPORT_TEMPLATE.md')
OUT_DOCX = os.path.join(BASE, 'MINI_PROJECT_REPORT.docx')

# Helper to set run font
def set_run_font(run, name='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = name
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_paragraph(doc, text, bold=False, underline=False, align=None, style=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', 12, bold=bold)
    if underline:
        run.font.underline = True
    if align:
        p.alignment = align
    if style:
        p.style = style
    return p


def insert_page_break(doc):
    doc.add_page_break()


def process_file_into_doc(doc, path):
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    # Insert content; insert page break when a line starts with 'PAGE ' or '---'
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            # blank line -> paragraph break
            doc.add_paragraph()
            continue
        # detect page separators used in our templates
        if stripped.startswith('PAGE') and '—' in stripped or stripped.startswith('PAGE'):
            # Insert a page break before this PAGE heading only if not at very start
            if doc.paragraphs:
                insert_page_break(doc)
            add_paragraph(doc, stripped, bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
            continue
        if stripped.startswith('CHAPTER') or stripped.startswith('CHAPTER'):
            # Chapter heading: uppercase, bold, centered
            add_paragraph(doc, stripped.upper(), bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
            continue
        # Headings in template start with numbers like '1.1' or '1.2' or SECTION labels - treat lines ending with ':' as bold headings
        if stripped.endswith(':') or stripped.endswith(')') and len(stripped) < 60:
            add_paragraph(doc, stripped, bold=True)
            continue
        # Regular content
        p = doc.add_paragraph(stripped)
        for run in p.runs:
            set_run_font(run, 'Times New Roman', 12)


if __name__ == '__main__':
    doc = Document()
    # Set default document font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Header and footer placeholders
    section = doc.sections[0]
    # Header: project title right-aligned
    header = section.header
    hdr_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hdr_p.text = '<PROJECT TITLE>'
    hdr_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    set_run_font(hdr_p.runs[0], 'Times New Roman', 12, bold=True)

    # Footer: page number placeholder bottom-left
    footer = section.footer
    f_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    f_p.text = 'Page: ' # User should insert page number field in Word
    f_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_run_font(f_p.runs[0], 'Times New Roman', 12)

    # Process front matter
    process_file_into_doc(doc, FRONT_MD)
    # ensure a page break before main content
    insert_page_break(doc)
    process_file_into_doc(doc, TEMPLATE_MD)

    doc.save(OUT_DOCX)
    print('Created', OUT_DOCX)
